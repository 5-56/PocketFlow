"""
Document Generator - 文档生成工具
支持生成Word、PDF、HTML等格式的文档
"""
import io
import base64
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import html

class DocumentGenerator:
    """文档生成器"""
    
    def __init__(self):
        self.default_font_size = 12
        self.default_font_name = "Arial"
        self.margin_inch = 1.0
        
    def generate_docx(self, content: Dict[str, Any], template_settings: Dict[str, Any], 
                     export_options: Dict[str, Any]) -> bytes:
        """生成Word文档"""
        try:
            # 创建文档
            doc = Document()
            
            # 应用模板设置
            self._apply_docx_template_settings(doc, template_settings)
            
            # 添加文档内容
            self._add_content_to_docx(doc, content, export_options)
            
            # 保存到字节流
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output.getvalue()
            
        except Exception as e:
            raise Exception(f"Word文档生成失败: {str(e)}")
    
    def generate_pdf(self, content: Dict[str, Any], template_settings: Dict[str, Any], 
                    export_options: Dict[str, Any]) -> bytes:
        """生成PDF文档"""
        try:
            output = io.BytesIO()
            
            # 页面设置
            page_size = A4
            if template_settings.get("page_size") == "letter":
                page_size = letter
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output,
                pagesize=page_size,
                leftMargin=self.margin_inch * inch,
                rightMargin=self.margin_inch * inch,
                topMargin=self.margin_inch * inch,
                bottomMargin=self.margin_inch * inch
            )
            
            # 准备内容
            story = []
            styles = getSampleStyleSheet()
            
            # 添加自定义样式
            self._add_pdf_styles(styles, template_settings)
            
            # 添加内容
            self._add_content_to_pdf(story, content, styles, export_options)
            
            # 构建PDF
            doc.build(story)
            
            return output.getvalue()
            
        except Exception as e:
            raise Exception(f"PDF文档生成失败: {str(e)}")
    
    def generate_html(self, content: Dict[str, Any], template_settings: Dict[str, Any], 
                     export_options: Dict[str, Any]) -> bytes:
        """生成HTML文档"""
        try:
            # HTML结构
            html_parts = []
            
            # 添加HTML头部
            html_parts.append(self._generate_html_head(template_settings))
            
            # 添加HTML主体
            html_parts.append("<body>")
            html_parts.append('<div class="document-container">')
            
            # 添加内容
            self._add_content_to_html(html_parts, content, export_options)
            
            html_parts.append("</div>")
            html_parts.append("</body>")
            html_parts.append("</html>")
            
            html_content = "\n".join(html_parts)
            
            return html_content.encode('utf-8')
            
        except Exception as e:
            raise Exception(f"HTML文档生成失败: {str(e)}")
    
    def _apply_docx_template_settings(self, doc: Document, settings: Dict[str, Any]):
        """应用Word文档模板设置"""
        try:
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(settings.get("top_margin", 1.0))
                section.bottom_margin = Inches(settings.get("bottom_margin", 1.0))
                section.left_margin = Inches(settings.get("left_margin", 1.0))
                section.right_margin = Inches(settings.get("right_margin", 1.0))
            
            # 设置默认字体
            font_name = settings.get("font_name", "Arial")
            font_size = settings.get("font_size", 12)
            
            # 修改Normal样式
            style = doc.styles['Normal']
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            
        except Exception as e:
            print(f"应用Word模板设置失败: {e}")
    
    def _add_content_to_docx(self, doc: Document, content: Dict[str, Any], options: Dict[str, Any]):
        """向Word文档添加内容"""
        try:
            # 添加标题（如果有）
            if options.get("add_title") and content.get("metadata", {}).get("title"):
                title = doc.add_heading(content["metadata"]["title"], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理文本元素
            for text_elem in content.get("text_elements", []):
                para = doc.add_paragraph()
                run = para.add_run(text_elem["content"])
                
                # 应用样式
                style = text_elem.get("style", {})
                if style.get("bold"):
                    run.bold = True
                if style.get("italic"):
                    run.italic = True
                if style.get("font_size"):
                    run.font.size = Pt(style["font_size"])
            
            # 处理图像元素
            for img_elem in content.get("image_elements", []):
                try:
                    # 添加图像
                    image_stream = io.BytesIO(img_elem["data"])
                    
                    # 计算图像尺寸
                    width = options.get("image_width", 6)  # 默认6英寸
                    doc.add_picture(image_stream, width=Inches(width))
                    
                    # 添加图像说明（如果有AI分析结果）
                    if options.get("add_image_captions") and img_elem.get("analysis"):
                        caption = doc.add_paragraph()
                        caption.add_run(f"图片说明: {img_elem['analysis'].get('description', '')}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                except Exception as e:
                    # 如果图像添加失败，添加错误说明
                    doc.add_paragraph(f"[图像显示失败: {str(e)}]")
            
            # 处理表格元素
            for table_elem in content.get("table_elements", []):
                table_data = table_elem["content"]
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            table.cell(i, j).text = str(cell_data)
            
        except Exception as e:
            raise Exception(f"添加Word文档内容失败: {str(e)}")
    
    def _add_pdf_styles(self, styles, settings: Dict[str, Any]):
        """添加PDF自定义样式"""
        try:
            # 标题样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                textColor=colors.black,
                alignment=1,  # 居中
                spaceAfter=30
            )
            styles.add(title_style)
            
            # 正文样式
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=settings.get("font_size", 12),
                textColor=colors.black,
                alignment=0,  # 左对齐
                spaceBefore=6,
                spaceAfter=6
            )
            styles.add(body_style)
            
        except Exception as e:
            print(f"添加PDF样式失败: {e}")
    
    def _add_content_to_pdf(self, story: List, content: Dict[str, Any], 
                           styles, options: Dict[str, Any]):
        """向PDF添加内容"""
        try:
            # 添加标题
            if options.get("add_title") and content.get("metadata", {}).get("title"):
                title = Paragraph(content["metadata"]["title"], styles['CustomTitle'])
                story.append(title)
                story.append(Spacer(1, 20))
            
            # 处理文本元素
            for text_elem in content.get("text_elements", []):
                para = Paragraph(html.escape(text_elem["content"]), styles['CustomBody'])
                story.append(para)
                story.append(Spacer(1, 6))
            
            # 处理图像元素
            for img_elem in content.get("image_elements", []):
                try:
                    image_stream = io.BytesIO(img_elem["data"])
                    
                    # 计算图像尺寸
                    img_width = options.get("image_width", 4) * inch
                    img = RLImage(image_stream, width=img_width)
                    
                    story.append(img)
                    story.append(Spacer(1, 12))
                    
                    # 添加图像说明
                    if options.get("add_image_captions") and img_elem.get("analysis"):
                        caption = Paragraph(
                            f"<i>图片说明: {html.escape(img_elem['analysis'].get('description', ''))}</i>",
                            styles['CustomBody']
                        )
                        story.append(caption)
                        story.append(Spacer(1, 12))
                    
                except Exception as e:
                    error_para = Paragraph(f"[图像显示失败: {str(e)}]", styles['CustomBody'])
                    story.append(error_para)
                    story.append(Spacer(1, 6))
            
            # 处理表格元素
            for table_elem in content.get("table_elements", []):
                table_data = table_elem["content"]
                if table_data:
                    # 创建表格
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    
                    story.append(table)
                    story.append(Spacer(1, 12))
            
        except Exception as e:
            raise Exception(f"添加PDF内容失败: {str(e)}")
    
    def _generate_html_head(self, settings: Dict[str, Any]) -> str:
        """生成HTML头部"""
        font_family = settings.get("font_name", "Arial, sans-serif")
        font_size = settings.get("font_size", 14)
        
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>智能文档处理系统 - 导出文档</title>
    <style>
        body {{
            font-family: {font_family};
            font-size: {font_size}px;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        
        .document-container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            border-radius: 8px;
        }}
        
        .title {{
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 30px;
            color: #333;
        }}
        
        .paragraph {{
            margin-bottom: 15px;
            text-align: justify;
            color: #444;
        }}
        
        .image-container {{
            text-align: center;
            margin: 20px 0;
        }}
        
        .image {{
            max-width: 100%;
            height: auto;
            border-radius: 4px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        
        .image-caption {{
            font-style: italic;
            color: #666;
            margin-top: 8px;
            font-size: 12px;
        }}
        
        .table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        
        .table th,
        .table td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        
        .table th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        
        .modified {{
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
            padding-left: 10px;
        }}
        
        .original {{
            opacity: 0.7;
        }}
    </style>
</head>"""
    
    def _add_content_to_html(self, html_parts: List[str], content: Dict[str, Any], 
                            options: Dict[str, Any]):
        """向HTML添加内容"""
        try:
            # 添加标题
            if options.get("add_title") and content.get("metadata", {}).get("title"):
                title = html.escape(content["metadata"]["title"])
                html_parts.append(f'<div class="title">{title}</div>')
            
            # 处理文本元素
            for text_elem in content.get("text_elements", []):
                css_class = "paragraph"
                if text_elem.get("modified"):
                    css_class += " modified"
                
                content_text = html.escape(text_elem["content"])
                html_parts.append(f'<div class="{css_class}">{content_text}</div>')
            
            # 处理图像元素
            for img_elem in content.get("image_elements", []):
                try:
                    # 将图像转换为base64
                    image_base64 = base64.b64encode(img_elem["data"]).decode('utf-8')
                    image_format = self._detect_image_format_from_data(img_elem["data"])
                    
                    html_parts.append('<div class="image-container">')
                    html_parts.append(
                        f'<img class="image" src="data:image/{image_format};base64,{image_base64}" alt="文档图片">'
                    )
                    
                    # 添加图像说明
                    if options.get("add_image_captions") and img_elem.get("analysis"):
                        caption = html.escape(img_elem["analysis"].get("description", ""))
                        html_parts.append(f'<div class="image-caption">图片说明: {caption}</div>')
                    
                    html_parts.append('</div>')
                    
                except Exception as e:
                    html_parts.append(f'<div class="paragraph">[图像显示失败: {html.escape(str(e))}]</div>')
            
            # 处理表格元素
            for table_elem in content.get("table_elements", []):
                table_data = table_elem["content"]
                if table_data:
                    html_parts.append('<table class="table">')
                    
                    for i, row_data in enumerate(table_data):
                        html_parts.append('<tr>')
                        tag = 'th' if i == 0 else 'td'
                        
                        for cell_data in row_data:
                            cell_content = html.escape(str(cell_data))
                            html_parts.append(f'<{tag}>{cell_content}</{tag}>')
                        
                        html_parts.append('</tr>')
                    
                    html_parts.append('</table>')
            
        except Exception as e:
            raise Exception(f"添加HTML内容失败: {str(e)}")
    
    def _detect_image_format_from_data(self, image_data: bytes) -> str:
        """从图像数据检测格式"""
        # 检查文件头
        if image_data.startswith(b'\xff\xd8\xff'):
            return 'jpeg'
        elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
            return 'png'
        elif image_data.startswith(b'BM'):
            return 'bmp'
        elif image_data.startswith(b'GIF'):
            return 'gif'
        elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:12]:
            return 'webp'
        else:
            return 'jpeg'  # 默认使用jpeg
    
    def create_template_settings(self, template_type: str = "default") -> Dict[str, Any]:
        """创建模板设置"""
        templates = {
            "default": {
                "font_name": "Arial",
                "font_size": 12,
                "top_margin": 1.0,
                "bottom_margin": 1.0,
                "left_margin": 1.0,
                "right_margin": 1.0,
                "page_size": "A4"
            },
            "professional": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "top_margin": 1.0,
                "bottom_margin": 1.0,
                "left_margin": 1.25,
                "right_margin": 1.25,
                "page_size": "A4"
            },
            "modern": {
                "font_name": "Calibri",
                "font_size": 11,
                "top_margin": 0.8,
                "bottom_margin": 0.8,
                "left_margin": 1.0,
                "right_margin": 1.0,
                "page_size": "A4"
            }
        }
        
        return templates.get(template_type, templates["default"])
    
    def create_export_options(self, options_type: str = "standard") -> Dict[str, Any]:
        """创建导出选项"""
        option_sets = {
            "standard": {
                "add_title": True,
                "add_image_captions": True,
                "image_width": 6,
                "include_original_content": False,
                "highlight_modifications": True
            },
            "minimal": {
                "add_title": False,
                "add_image_captions": False,
                "image_width": 4,
                "include_original_content": False,
                "highlight_modifications": False
            },
            "detailed": {
                "add_title": True,
                "add_image_captions": True,
                "image_width": 6,
                "include_original_content": True,
                "highlight_modifications": True
            }
        }
        
        return option_sets.get(options_type, option_sets["standard"])