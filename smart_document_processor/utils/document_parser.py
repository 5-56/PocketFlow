"""
Document Parser - 文档解析工具
支持Word、PDF、文本文件等格式的解析
"""
import io
import base64
import zipfile
from typing import Dict, List, Any, Optional
from docx import Document
import PyPDF2
from PIL import Image
import xml.etree.ElementTree as ET

class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_formats = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/pdf': 'pdf',
            'text/plain': 'txt',
            'text/html': 'html',
            'application/msword': 'doc'
        }
    
    def parse(self, file_data: bytes, file_name: str, content_type: str) -> Dict[str, Any]:
        """解析文档"""
        
        # 确定文件格式
        file_format = self._detect_format(content_type, file_name)
        
        if file_format == 'docx':
            return self._parse_docx(file_data, file_name)
        elif file_format == 'pdf':
            return self._parse_pdf(file_data, file_name)
        elif file_format == 'txt':
            return self._parse_text(file_data, file_name)
        elif file_format == 'html':
            return self._parse_html(file_data, file_name)
        else:
            raise ValueError(f"不支持的文件格式: {file_format}")
    
    def _detect_format(self, content_type: str, file_name: str) -> str:
        """检测文件格式"""
        if content_type in self.supported_formats:
            return self.supported_formats[content_type]
        
        # 根据文件扩展名检测
        if file_name.lower().endswith('.docx'):
            return 'docx'
        elif file_name.lower().endswith('.pdf'):
            return 'pdf'
        elif file_name.lower().endswith('.txt'):
            return 'txt'
        elif file_name.lower().endswith('.html'):
            return 'html'
        else:
            return 'unknown'
    
    def _parse_docx(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析Word文档"""
        try:
            # 创建Document对象
            doc = Document(io.BytesIO(file_data))
            
            elements = []
            
            # 解析段落
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    elements.append({
                        "type": "text",
                        "content": paragraph.text,
                        "position": {"paragraph": i},
                        "style": {
                            "font_size": self._get_paragraph_font_size(paragraph),
                            "bold": self._is_paragraph_bold(paragraph),
                            "italic": self._is_paragraph_italic(paragraph)
                        }
                    })
            
            # 解析图片
            relationships = doc.part.rels
            for rel_id, rel in relationships.items():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        elements.append({
                            "type": "image",
                            "data": image_data,
                            "position": {"rel_id": rel_id},
                            "size": self._get_image_size(image_data)
                        })
                    except Exception as e:
                        print(f"图片解析失败: {e}")
            
            # 解析表格
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                elements.append({
                    "type": "table",
                    "content": table_data,
                    "position": {"table": i}
                })
            
            return {
                "file_name": file_name,
                "file_format": "docx",
                "elements": elements,
                "metadata": {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "page_count": len(doc.sections)
                }
            }
            
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")
    
    def _parse_pdf(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析PDF文档"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            
            elements = []
            
            # 解析每一页
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    # 按段落分割文本
                    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                    
                    for i, paragraph in enumerate(paragraphs):
                        elements.append({
                            "type": "text",
                            "content": paragraph,
                            "position": {"page": page_num, "paragraph": i},
                            "style": {}
                        })
            
            return {
                "file_name": file_name,
                "file_format": "pdf",
                "elements": elements,
                "metadata": {
                    "page_count": len(pdf_reader.pages),
                    "text_extracted": True
                }
            }
            
        except Exception as e:
            raise Exception(f"PDF文档解析失败: {str(e)}")
    
    def _parse_text(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析纯文本文档"""
        try:
            text_content = file_data.decode('utf-8')
            
            elements = []
            
            # 按段落分割
            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
            
            for i, paragraph in enumerate(paragraphs):
                elements.append({
                    "type": "text",
                    "content": paragraph,
                    "position": {"paragraph": i},
                    "style": {}
                })
            
            return {
                "file_name": file_name,
                "file_format": "txt",
                "elements": elements,
                "metadata": {
                    "character_count": len(text_content),
                    "line_count": len(text_content.split('\n'))
                }
            }
            
        except Exception as e:
            raise Exception(f"文本文档解析失败: {str(e)}")
    
    def _parse_html(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析HTML文档"""
        try:
            html_content = file_data.decode('utf-8')
            
            # 简单的HTML解析，移除标签
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            
            elements = []
            
            # 按段落分割
            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph:
                    elements.append({
                        "type": "text",
                        "content": paragraph,
                        "position": {"paragraph": i},
                        "style": {}
                    })
            
            return {
                "file_name": file_name,
                "file_format": "html",
                "elements": elements,
                "metadata": {
                    "original_html": True
                }
            }
            
        except Exception as e:
            raise Exception(f"HTML文档解析失败: {str(e)}")
    
    def _get_paragraph_font_size(self, paragraph) -> Optional[int]:
        """获取段落字体大小"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].font.size.pt if paragraph.runs[0].font.size else None
        except:
            pass
        return None
    
    def _is_paragraph_bold(self, paragraph) -> bool:
        """检查段落是否加粗"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].bold or False
        except:
            pass
        return False
    
    def _is_paragraph_italic(self, paragraph) -> bool:
        """检查段落是否斜体"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].italic or False
        except:
            pass
        return False
    
    def _get_image_size(self, image_data: bytes) -> Dict[str, int]:
        """获取图片尺寸"""
        try:
            image = Image.open(io.BytesIO(image_data))
            return {"width": image.width, "height": image.height}
        except:
            return {"width": 0, "height": 0}
    
    def extract_images(self, file_data: bytes, file_format: str) -> List[Dict[str, Any]]:
        """单独提取图片"""
        images = []
        
        if file_format == 'docx':
            try:
                doc = Document(io.BytesIO(file_data))
                relationships = doc.part.rels
                
                for rel_id, rel in relationships.items():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            size = self._get_image_size(image_data)
                            images.append({
                                "id": rel_id,
                                "data": image_data,
                                "size": size,
                                "format": self._detect_image_format(image_data)
                            })
                        except Exception as e:
                            print(f"图片提取失败: {e}")
            except Exception as e:
                print(f"Word文档图片提取失败: {e}")
        
        return images
    
    def _detect_image_format(self, image_data: bytes) -> str:
        """检测图片格式"""
        try:
            image = Image.open(io.BytesIO(image_data))
            return image.format.lower()
        except:
            return "unknown"